from typing import Any

from .base import BaseParser


class DocxParser(BaseParser):
    def _extract(self, file_path: str) -> dict[str, Any]:
        import docx
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            "text": text,
            "pages": [{"page_num": 1, "text": text}],
            "tables": tables,
            "images": [],
            "metadata": {"source": file_path},
            "sections": []
        }
